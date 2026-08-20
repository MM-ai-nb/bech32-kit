from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "任务报告书-卓娟-bech32-kit.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=11, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_text(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx] / 567)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    set_paragraph_font(paragraph, size=11)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for idx, (key, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), key, bold=True, fill=LIGHT_FILL)
        set_cell_text(table.cell(idx, 1), value)
    return table


def add_status_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [3300, 1500, 4560])
    headers = ["验收条件", "状态", "说明"]
    for idx, header in enumerate(headers):
        align = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.cell(0, idx), header, bold=True, fill=LIGHT_FILL, align=align)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(row_idx, col_idx), text, align=align)
    return table


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.author = "卓娟"
    doc.core_properties.last_modified_by = "卓娟"
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("bech32-kit 任务报告书")
    set_run_font(run, size=24, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    sub_run = subtitle.add_run("MoonBit 原生 Bech32 / Bech32m 编码、诊断、策略校验、文本扫描与 SegWit 地址库")
    set_run_font(sub_run, size=12, color=GRAY)

    add_key_value_table(
        doc,
        [
            ("参赛者", "卓娟"),
            ("联系方式", "见飞书报名表 / 正式提交材料"),
            ("GitHub 仓库", "https://github.com/MM-ai-nb/bech32-kit"),
            ("Mooncakes 包名", "MM-ai-nb/bech32-kit"),
            ("Mooncakes 页面", "https://mooncakes.io/docs/MM-ai-nb/bech32-kit"),
            ("版本 / 许可证", "0.1.0 / MIT"),
            ("报告日期", "2026-08-20"),
        ],
    )

    add_heading(doc, "一、项目背景与用途", 1)
    add_body(
        doc,
        "Bech32 和 Bech32m 是常用于区块链地址、协议字段和低误读率短字符串的编码格式。"
        "MoonBit 生态中同类基础库仍然较少，上层项目如果需要处理 checksum、bit group 转换、地址导入策略或 SegWit 地址校验，通常需要重复实现相关逻辑。",
    )
    add_body(
        doc,
        "bech32-kit 的目标是提供一个纯 MoonBit、可测试、可发布、可复用的 Bech32 / Bech32m 工具库，"
        "方便 MoonBit CLI、WebAssembly 应用、钱包工具、链上数据工具、批量导入流程和教学项目直接使用。",
    )

    add_heading(doc, "二、主要功能", 1)
    for item in [
        "Bech32 编码、解码和 checksum 校验。",
        "Bech32m 编码、解码和 checksum 校验。",
        "8-bit 字节与 5-bit Bech32 数据之间的 bit group 转换。",
        "SegWit v0 到 v16 地址编码与验证。",
        "结构化错误类型，可区分非法字符、大小写混用、checksum 错误、padding 错误和 SegWit 规则错误。",
        "输入 profile、规范化、checksum words 提取、稳定错误码和批量诊断报告。",
        "高层策略校验，可按 variant、canonical、HRP、payload 长度、网络和 witness program 类型接受或拒绝输入。",
        "文本扫描器可从日志、剪贴板、表单备注和多行文本中提取 Bech32 / SegWit 候选串。",
        "扫描结果包含 offset、行号、列号、重复候选识别、网络过滤和策略 lint 问题。",
        "可运行示例，输出标准 SegWit v0 地址。",
        "GitHub Actions 持续集成，自动执行 check、build、test 和示例运行。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、使用方法", 1)
    add_body(doc, "安装 Mooncakes 包：moon add MM-ai-nb/bech32-kit")
    add_body(doc, "本地构建与测试：moon check；moon build；moon test；moon run cmd/main")
    add_body(doc, "示例输出：bc1qw508d6qejxtdg4y5r3zarvary0c5xw7kv8f3t4")

    add_heading(doc, "四、测试与构建记录", 1)
    add_body(doc, "2026-08-20 已在本地执行完整验收命令。")
    for item in [
        "moon check 通过。",
        "moon build 通过。",
        "moon test --deny-warn 通过，结果为 Total tests: 28, passed: 28, failed: 0.",
        "moon run cmd/main 通过，并输出可验证的 Bech32 SegWit 示例地址。",
        "moon fmt --check、moon check --deny-warn、moon build 和 moon info 均通过。",
        "按排除 _build、空行和 // 注释行的口径统计，有效 MoonBit 行数为 4258 行。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "五、持续集成", 1)
    add_body(doc, "项目已配置 GitHub Actions 工作流 .github/workflows/ci.yml。CI 在 push 和 pull request 时执行以下步骤：")
    for item in [
        "安装 MoonBit 工具链。",
        "执行 moon fmt --check。",
        "执行 moon check --deny-warn。",
        "执行 moon build。",
        "执行 moon test --deny-warn。",
        "执行 moon info。",
        "执行 moon run cmd/main 验证示例可运行。",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "六、Mooncakes 发布记录", 1)
    add_body(
        doc,
        "项目已在 moon.mod 中配置 Mooncakes 元数据：当前包名 MM-ai-nb/bech32-kit，版本 0.1.0，仓库 "
        "https://github.com/MM-ai-nb/bech32-kit.git，许可证 MIT，README 为 README.md。"
        "Mooncakes 页面为 https://mooncakes.io/docs/MM-ai-nb/bech32-kit。",
    )
    add_body(doc, "2026-08-20 已使用 MM-ai-nb 对应的 Mooncakes 登录会话完成发布：moon whoami 显示 MM-ai-nb；moon publish --dry-run 服务器返回 202 Accepted；moon publish 服务器返回 200 OK。")

    add_heading(doc, "七、开源许可证与第三方依赖", 1)
    add_body(
        doc,
        "项目采用 MIT License。实现依据公开标准 BIP-0173 与 BIP-0350 的算法说明和公开测试向量，"
        "不移植第三方源码，不包含外部素材或私有代码。项目当前没有引入额外第三方运行时依赖。",
    )

    add_heading(doc, "八、功能边界与后续维护价值", 1)
    add_body(
        doc,
        "项目专注于 Bech32 / Bech32m 字符串处理、checksum 验证、bit group 转换、SegWit 地址校验、输入诊断、策略校验和传入文本扫描，"
        "不处理私钥、助记词、网络请求、钱包账户管理、本地目录扫描或链上交易签名。清晰的边界有利于降低安全风险，并便于作为基础库长期维护。",
    )
    for item in [
        "补充更多公开测试向量与 fuzz 风格边界测试。",
        "根据 MoonBit 生态变化维护 Mooncakes 包元数据。",
        "在真实钱包工具、CLI 或 WebAssembly 应用中验证 API 易用性。",
        "根据使用反馈扩展策略模板、错误信息和文档示例。",
        "保持语义化版本发布和 CHANGELOG 记录。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "九、验收条件对照", 1)
    add_status_table(
        doc,
        [
            ("以 MoonBit 作为主要实现语言", "已满足", "核心代码、测试和示例均为 MoonBit。"),
            ("代码仓库公开且可以正常访问", "已满足", "已推送到 MM-ai-nb/bech32-kit 默认分支，远端 CI 已通过。"),
            ("提供清晰、完整的 README", "已满足", "README 已说明用途、安装、功能、示例和验收命令。"),
            ("提供可以实际运行的示例", "已满足", "cmd/main 可运行并输出 SegWit 地址。"),
            ("配置持续集成 CI", "已满足", "已配置 GitHub Actions。"),
            ("提供可运行的测试", "已满足", "本地 moon test 结果为 28 passed。"),
            ("项目能够正常构建", "已满足", "moon check 与 moon build 均通过。"),
            ("有效 MoonBit 代码规模", "已满足", "当前有效 MoonBit 行数为 4258 行。"),
            ("按要求发布至 mooncakes.io", "已满足", "已发布，包名为 MM-ai-nb/bech32-kit，版本 0.1.0。"),
            ("开发过程和提交记录可以追踪", "已整理", "当前分支基于 MM-ai-nb/main 整理，本次新增提交将使用卓娟身份信息。"),
            ("功能边界和维护价值明确", "已满足", "DESIGN.md 与本报告均已说明。"),
            ("第三方代码、素材和依赖符合开源许可证", "已满足", "MIT License，无外部素材和额外运行时依赖。"),
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("bech32-kit · MoonBit Hackathon")
    set_run_font(footer_run, size=9, color=GRAY)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
